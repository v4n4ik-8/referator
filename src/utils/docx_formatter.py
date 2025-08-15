from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from models import Essay, Section

class DocumentFormatter:
    def __init__(self):
        self.font_name = 'Times New Roman'
        self.font_size = 14
    
    def create_document(self, essay: Essay) -> Document:
        """Создает отформатированный документ из реферата"""
        doc = Document()
        
        # Установка стиля для всего документа
        style = doc.styles['Normal']
        font = style.font
        font.name = self.font_name
        font.size = Pt(self.font_size)
        
        # Добавление заголовка
        title = doc.add_paragraph()
        title_run = title.add_run(essay.topic)
        title_run.bold = True
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Добавление содержимого
        for section in essay.sections:
            # Пропускаем заключение
            if "Заключение" in section.title:
                continue
                
            # Добавляем заголовок раздела
            heading = doc.add_paragraph()
            heading_run = heading.add_run(section.title)
            heading_run.bold = True
            
            # Устанавливаем выравнивание заголовка
            if "Введение" in section.title:
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Очищаем markdown и удаляем дублирование заголовка
            clean_content = self._clean_markdown(section.content)
            clean_content = self._remove_duplicate_title(section.title, clean_content)
            
            # Добавляем содержимое раздела с выравниванием по ширине
            content_paragraph = doc.add_paragraph(clean_content)
            content_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        # Добавляем нумерацию страниц
        self._add_page_numbers(doc)
        
        return doc
    
    def _clean_markdown(self, text: str) -> str:
        """Удаляет markdown-форматирование из текста"""
        # Удаляем заголовки (#)
        lines = []
        for line in text.split('\n'):
            if line.strip().startswith('#'):
                line = line.lstrip('#').strip()
            lines.append(line)
        
        text = '\n'.join(lines)
        
        # Удаляем звездочки для жирного и курсивного текста
        text = text.replace('**', '').replace('*', '')
        
        # Удаляем бэктики
        text = text.replace('`', '')
        
        # Удаляем маркеры списков
        text = text.replace('- ', '').replace('* ', '')
        
        return text
    
    def _remove_duplicate_title(self, title: str, content: str) -> str:
        """Удаляет дублирование заголовка из начала контента"""
        # Очищаем заголовок и контент для сравнения
        clean_title = title.strip().lower()
        content_lines = content.strip().split('\n')
        
        # Проверяем, начинается ли контент с заголовка
        if content_lines and content_lines[0].strip().lower() == clean_title:
            # Удаляем первую строку и все пустые строки после неё
            while content_lines and not content_lines[0].strip():
                content_lines.pop(0)
            return '\n'.join(content_lines).strip()
        
        return content.strip()
    
    def _add_page_numbers(self, doc: Document):
        """Добавляет нумерацию страниц внизу (Простой пример 2)"""
        section = doc.sections[0]
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        field = OxmlElement('w:fldSimple')
        field.set(qn('w:instr'), 'PAGE')
        run = paragraph.add_run()
        run._r.append(field)
        run.font.name = self.font_name
        run.font.size = Pt(self.font_size)
        
    
    def add_table_of_contents(self, doc: Document) -> None:
        """Добавляет оглавление"""
        # TODO: Реализовать добавление оглавления
        pass
    
    def add_title_page(self, doc: Document, essay: Essay) -> None:
        """Добавляет титульный лист"""
        # TODO: Реализовать добавление титульного листа по шаблону
        pass 